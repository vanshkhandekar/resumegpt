#!/usr/bin/env python3
"""
AI Resume Studio — Complete University-Level Thesis Generator
Generates ~80 page PDF with Times New Roman, proper formatting, diagrams
"""

import os
import sys
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch, mm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.colors import HexColor, black
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle,
    Preformatted, KeepTogether, ListFlowable, ListItem, Image
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Try to register Times New Roman
FONT_NAME = "Times-Roman"
FONT_BOLD = "Times-Bold"
FONT_ITALIC = "Times-Italic"
FONT_BI = "Times-BoldItalic"

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "AI_Resume_Studio_Thesis.pdf")

WIDTH, HEIGHT = A4
MARGIN_LEFT = 1.5 * inch
MARGIN_RIGHT = 1.0 * inch
MARGIN_TOP = 1.0 * inch
MARGIN_BOTTOM = 1.0 * inch

# ─── Styles ───
def get_styles():
    styles = getSampleStyleSheet()
    
    # Based on Notice: Heading=16, Sub-Heading=14, Normal=12
    styles.add(ParagraphStyle(
        'ThesisTitle', fontName=FONT_BOLD, fontSize=22,
        leading=28, alignment=TA_CENTER, spaceAfter=20
    ))
    styles.add(ParagraphStyle(
        'ChapterTitle', fontName=FONT_BOLD, fontSize=16,
        leading=26, spaceAfter=40, spaceBefore=40,
        textColor=HexColor('#000000'), alignment=TA_CENTER
    ))
    styles.add(ParagraphStyle(
        'SectionTitle', fontName=FONT_BOLD, fontSize=14,
        leading=24, spaceAfter=24, spaceBefore=30,
        textColor=HexColor('#000000')
    ))
    styles.add(ParagraphStyle(
        'SubSection', fontName=FONT_BOLD, fontSize=14, # Instructions say Sub-Heading 14
        leading=22, spaceAfter=20, spaceBefore=24
    ))
    styles.add(ParagraphStyle(
        'Body', fontName=FONT_NAME, fontSize=12,
        leading=22, alignment=TA_JUSTIFY, spaceAfter=20 # Even more spacing
    ))
    styles.add(ParagraphStyle(
        'BodyIndent', fontName=FONT_NAME, fontSize=12,
        leading=22, alignment=TA_JUSTIFY, spaceAfter=20,
        leftIndent=24
    ))
    styles.add(ParagraphStyle(
        'ThesisBullet', fontName=FONT_NAME, fontSize=12,
        leading=20, alignment=TA_JUSTIFY, spaceAfter=10,
        leftIndent=36, bulletIndent=18
    ))
    styles.add(ParagraphStyle(
        'CenterNormal', fontName=FONT_NAME, fontSize=12,
        leading=20, alignment=TA_CENTER, spaceAfter=10
    ))
    styles.add(ParagraphStyle(
        'CenterBold', fontName=FONT_BOLD, fontSize=12,
        leading=20, alignment=TA_CENTER, spaceAfter=10
    ))
    styles.add(ParagraphStyle(
        'CodeBlock', fontName='Courier', fontSize=11, # Increased font size
        leading=14, spaceAfter=12, spaceBefore=10,
        leftIndent=24, rightIndent=12,
        backColor=HexColor('#f5f5f5')
    ))
    styles.add(ParagraphStyle(
        'Caption', fontName=FONT_ITALIC, fontSize=10,
        leading=14, alignment=TA_CENTER, spaceAfter=12
    ))
    styles.add(ParagraphStyle(
        'Reference', fontName=FONT_NAME, fontSize=11,
        leading=16, spaceAfter=8, leftIndent=36,
        firstLineIndent=-36
    ))
    styles.add(ParagraphStyle(
        'TOCEntry', fontName=FONT_NAME, fontSize=12,
        leading=20, spaceAfter=2
    ))
    styles.add(ParagraphStyle(
        'TOCChapter', fontName=FONT_BOLD, fontSize=12,
        leading=22, spaceAfter=2
    ))
    styles.add(ParagraphStyle(
        'ThesisSubtitle', fontName=FONT_ITALIC, fontSize=14,
        leading=20, alignment=TA_CENTER, spaceAfter=8
    ))
    return styles

def add_header_footer(canvas, doc):
    page_num = canvas.getPageNumber()
    canvas.saveState()
    canvas.setFont(FONT_NAME, 10)
    
    if page_num > 6:  # After starting pages and Index
        # Header: Name of Project (Right) - Project Name is AI Resume Studio
        canvas.drawRightString(WIDTH - MARGIN_RIGHT, HEIGHT - 0.5*inch, "AI Resume Studio")
        
        # Footer: College Name (Left)
        canvas.drawString(MARGIN_LEFT, 0.5*inch, "JANAPRABHA COLLEGE, RAMTEK")
        
        # Footer: Page Number (Right)
        canvas.drawRightString(WIDTH - MARGIN_RIGHT, 0.5*inch, f"{page_num}")
    else:
        # For starting pages, maybe only page number or nothing
        pass
        
    canvas.restoreState()

# ─── Content Sections ───
from thesis_content import (
    build_front_matter, build_toc, build_chapter1,
    build_chapter2, build_chapter3, build_chapter4,
    build_chapter5, build_chapter6, build_references,
    build_appendix
)

def main():
    doc = SimpleDocTemplate(
        OUTPUT_FILE, pagesize=A4,
        leftMargin=MARGIN_LEFT, rightMargin=MARGIN_RIGHT,
        topMargin=MARGIN_TOP, bottomMargin=MARGIN_BOTTOM,
        title="AI Resume Studio with ATS & AI Features — MCA Thesis",
        author="Vansh Khandekar"
    )
    
    styles = get_styles()
    story = []
    
    print("🔨 Building all sections...")
    story += build_front_matter(styles)
    # build_front_matter already ends with a PageBreak()
    
    story += build_toc(styles)
    # build_toc already ends with a PageBreak()
    
    story += build_chapter1(styles)
    story.append(Spacer(1, 40))
    
    story += build_chapter2(styles)
    story.append(Spacer(1, 40))
    
    story += build_chapter3(styles)
    story.append(Spacer(1, 40))
    
    story += build_chapter4(styles)
    story.append(Spacer(1, 40))
    
    story += build_chapter5(styles)
    story.append(Spacer(1, 40))
    
    story += build_chapter6(styles)
    story.append(Spacer(1, 40))
    
    story += build_references(styles)
    story.append(Spacer(1, 40))
    
    story += build_appendix(styles)
    
    print(f"\n📝 Generating PDF: {OUTPUT_FILE}")
    story_copy = list(story) # Copy before build() consumes it
    doc.build(story, onFirstPage=add_header_footer, onLaterPages=add_header_footer)
    print(f"✅ Thesis PDF generated successfully!")

    # ─── DOCX Generation ───
    print(f"📄 Generating DOCX: {OUTPUT_FILE.replace('.pdf', '.docx')}")
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        word_doc = Document()
        
        for flowable in story_copy:
            if isinstance(flowable, Paragraph):
                p = word_doc.add_paragraph()
                text = flowable.getPlainText()
                style_name = flowable.style.name
                
                # Simple style mapping
                run = p.add_run(text)
                if 'Title' in style_name or 'Chapter' in style_name:
                    run.bold = True
                    run.font.size = Pt(16)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif 'Section' in style_name:
                    run.bold = True
                    run.font.size = Pt(14)
                elif 'Caption' in style_name:
                    run.italic = True
                    run.font.size = Pt(10)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    run.font.size = Pt(12)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
            elif isinstance(flowable, Preformatted):
                p = word_doc.add_paragraph()
                # Preformatted also works with getPlainText() or flowable.text sometimes
                # but to be safe use getPlainText()
                try: text = flowable.getPlainText()
                except: text = str(flowable)
                run = p.add_run(text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                
            elif isinstance(flowable, Image):
                try:
                    word_doc.add_picture(flowable.filename, width=Inches(6.0))
                except:
                    pass
                    
            elif isinstance(flowable, Spacer):
                word_doc.add_paragraph()
                
            elif isinstance(flowable, PageBreak):
                word_doc.add_page_break()
                
            elif isinstance(flowable, Table):
                # Simple table conversion
                t_data = flowable._cellvalues
                word_table = word_doc.add_table(rows=len(t_data), cols=len(t_data[0]))
                word_table.style = 'Table Grid'
                for r_idx, row in enumerate(t_data):
                    for c_idx, val in enumerate(row):
                        word_table.cell(r_idx, c_idx).text = str(val)

        word_doc.save(OUTPUT_FILE.replace('.pdf', '.docx'))
        print(f"✅ Thesis DOCX generated successfully!")
    except Exception as e:
        print(f"❌ DOCX generation failed: {e}")

    print(f"📄 Location: {OUTPUT_DIR}")

if __name__ == "__main__":
    main()
