import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, Image, Preformatted
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT

# --- SHARED DATA ---
TOC_DATA = [
    ("Sr. No.", "CHAPTER", "Page No"),
    ("1", "BRIEF REVIEW OF THE PROJECT", ""),
    ("", "1.1 TITLE", "1"),
    ("", "1.2 INTRODUCTION / OBJECTIVE", "2"),
    ("", "1.3 PRELIMINARY INVESTIGATION", "3"),
    ("", "1.4 FLAWS IN PRESENT SYSTEM", "4"),
    ("", "1.5 NEED OF NEW SYSTEM", "5"),
    ("2", "DETAILED SYSTEM DESIGN", ""),
    ("", "2.1 SYSTEM FLOW CHART", "6"),
    ("", "2.2 DETAILS OF LOGIC DEVELOPED", "7"),
    ("", "2.3 STRUCTURE DIAGRAM OF EACH MODULE", "9"),
    ("", "2.4 DATA DICTIONARY", "10"),
    ("", "2.5 DATA FLOW DIAGRAMS", "12"),
    ("3", "SOFTWARE/ HARDWARE DETAILS", ""),
    ("", "3.1 CHOICE OF A LANGUAGE USED", "14"),
    ("", "3.2 HARDWARE/ SOFTWARE SPECIFICATION", "15"),
    ("4", "SYSTEM DESIGN", ""),
    ("", "4.1 PROGRAM LISTING", "16"),
    ("", "4.2 INPUT SCREENS", "18"),
    ("", "4.3 OUTPUT SCREENS / REPORTS", "22"),
    ("5", "USER DOCUMENTATION", ""),
    ("", "5.1 IMPLEMENTATION, PROGRAM\nEXECUTION & MAINTENANCE", "26"),
    ("6", "CONCLUSION", ""),
    ("", "6.1 LIMITATIONS OF THE SYSTEM", "28"),
    ("", "6.2 SCOPE AND FUTURE MODIFICATION", "29"),
    ("7", "REFERENCES / BIBLIOGRAPHY", "30")
]

# --- DOCX GENERATION ---
def setup_docx_hf(doc):
    for section in doc.sections:
        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = "AI Resume Studio\n" + "—" * 60
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.text = "—" * 60 + "\nJanaprabha Institute of Engineering and Technology\t\t\t\tPg. "
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_docx_heading(doc, text, level=1, align=WD_ALIGN_PARAGRAPH.LEFT):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.bold = True
    if level == 2: run.font.underline = True
    h.alignment = align
    return h

def add_docx_para(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.alignment = align
    return p

def add_docx_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def generate_docx(filepath):
    doc = Document()
    setup_docx_hf(doc)
    
    add_docx_heading(doc, 'CONTENTS', level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    table = doc.add_table(rows=len(TOC_DATA), cols=3)
    table.style = 'Table Grid'
    for i, (sr, chap, pg) in enumerate(TOC_DATA):
        cells = table.rows[i].cells
        cells[0].text = sr
        cells[1].text = chap
        cells[2].text = pg
        for cell in cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    if sr != "" or i == 0: r.font.bold = True

    # Ch 1
    doc.add_page_break()
    add_docx_heading(doc, 'BRIEF REVIEW OF THE PROJECT', level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_docx_heading(doc, '1.1 TITLE', level=2)
    add_docx_para(doc, "AI Resume Studio — Resume Maker with ATS Scoring & AI-Powered Features")
    add_docx_heading(doc, '1.2 INTRODUCTION / OBJECTIVE', level=2)
    add_docx_para(doc, "The online resume building ecosystem is a very popular platform for job seekers. Our AI Resume Studio provides facilities like smart skill detection, auto ATS scores, and intelligent summaries.")
    
    # Diagrams
    doc.add_page_break()
    add_docx_heading(doc, 'DETAILED SYSTEM DESIGN', level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_docx_heading(doc, '2.1 SYSTEM FLOW CHART', level=2)
    if os.path.exists('data/diagrams/system_flow_admin.png'):
        doc.add_picture('data/diagrams/system_flow_admin.png', width=Inches(4.5))
    if os.path.exists('data/diagrams/system_flow_user.png'):
        doc.add_picture('data/diagrams/system_flow_user.png', width=Inches(3.5))

    add_docx_heading(doc, '2.3 STRUCTURE DIAGRAM OF EACH MODULE', level=2)
    if os.path.exists('data/diagrams/activity_admin.png'):
        doc.add_picture('data/diagrams/activity_admin.png', width=Inches(5))
    if os.path.exists('data/diagrams/activity_user.png'):
        doc.add_picture('data/diagrams/activity_user.png', width=Inches(5))
        
    # Code
    doc.add_page_break()
    add_docx_heading(doc, '4.1 PROGRAM LISTING', level=2)
    try:
        with open('src/pages/Index.tsx', 'r') as f:
            code = f.read()[:1200] + "\n... [Truncated]"
        p = doc.add_paragraph()
        r = p.add_run(code)
        r.font.name = 'Courier New'
        r.font.size = Pt(8)
    except: pass

    # Screenshots
    doc.add_page_break()
    add_docx_heading(doc, '4.2 INPUT SCREENS', level=2)
    if os.path.exists('public/screenshots/real_landing.png'):
        doc.add_picture('public/screenshots/real_landing.png', width=Inches(6))
    
    doc.add_page_break()
    add_docx_heading(doc, '4.3 OUTPUT SCREENS / REPORTS', level=2)
    if os.path.exists('public/screenshots/real_builder.png'):
        doc.add_picture('public/screenshots/real_builder.png', width=Inches(6))
    if os.path.exists('public/screenshots/real_dashboard.png'):
        doc.add_picture('public/screenshots/real_dashboard.png', width=Inches(6))

    doc.save(filepath)

# --- PDF GENERATION (REPORTLAB) ---
def add_pdf_hf(canvas, doc):
    canvas.saveState()
    canvas.setFont('Times-Bold', 10)
    canvas.drawRightString(A4[0]-inch, A4[1]-0.5*inch, "AI Resume Studio")
    canvas.line(inch, A4[1]-0.65*inch, A4[0]-inch, A4[1]-0.65*inch)
    
    canvas.line(inch, 0.75*inch, A4[0]-inch, 0.75*inch)
    canvas.drawString(inch, 0.6*inch, "Janaprabha Institute of Engineering and Technology")
    canvas.drawRightString(A4[0]-inch, 0.6*inch, f"Pg. {doc.page}")
    canvas.restoreState()

def generate_pdf(filepath):
    doc = SimpleDocTemplate(filepath, pagesize=A4, rightMargin=inch, leftMargin=inch, topMargin=inch, bottomMargin=inch)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle('BcaCenterH1', parent=styles['Heading1'], alignment=TA_CENTER, fontName='Times-Bold', fontSize=18))
    styles.add(ParagraphStyle('BcaH2', parent=styles['Heading2'], fontName='Times-Bold', fontSize=14, spaceBefore=12, underline=True))
    styles.add(ParagraphStyle('BcaBodyText', parent=styles['Normal'], fontName='Times-Roman', fontSize=12, leading=14, alignment=TA_JUSTIFY))
    styles.add(ParagraphStyle('BcaCode', parent=styles['Normal'], fontName='Courier', fontSize=8, leading=10, backColor=colors.lightgrey, leftIndent=20))

    story = []
    
    # TOC
    story.append(Paragraph("CONTENTS", styles['BcaCenterH1']))
    story.append(Spacer(1, 0.2*inch))
    t = Table(TOC_DATA, colWidths=[0.8*inch, 4*inch, 1*inch])
    t.setStyle(TableStyle([
        ('GRID', (0,0), (-1,-1), 1, colors.black),
        ('FONTNAME', (0,0), (-1,-1), 'Times-Roman'),
        ('FONTNAME', (0,0), (0,-1), 'Times-Bold'), # Sr No Bold
        ('FONTNAME', (0,0), (-1,0), 'Times-Bold'), # Header Bold
    ]))
    story.append(t)
    story.append(PageBreak())

    # Ch 1
    story.append(Paragraph("BRIEF REVIEW OF THE PROJECT", styles['BcaCenterH1']))
    story.append(Paragraph("1.1 TITLE", styles['BcaH2']))
    story.append(Paragraph("AI Resume Studio — Resume Maker with ATS Scoring & AI-Powered Features", styles['BcaBodyText']))
    story.append(Paragraph("1.2 INTRODUCTION / OBJECTIVE", styles['BcaH2']))
    story.append(Paragraph("The online resume building ecosystem is a very popular platform for job seekers. Our AI Resume Studio provides facilities like smart skill detection, auto ATS scores, and intelligent summaries.", styles['BcaBodyText']))
    
    # Diagrams
    story.append(PageBreak())
    story.append(Paragraph("DETAILED SYSTEM DESIGN", styles['BcaCenterH1']))
    story.append(Paragraph("2.1 SYSTEM FLOW CHART", styles['BcaH2']))
    if os.path.exists('data/diagrams/system_flow_admin.png'):
        story.append(Image('data/diagrams/system_flow_admin.png', width=4.5*inch, height=3*inch))
    story.append(Spacer(1, 0.2*inch))
    if os.path.exists('data/diagrams/system_flow_user.png'):
        story.append(Image('data/diagrams/system_flow_user.png', width=3.5*inch, height=4*inch))

    story.append(Paragraph("2.3 STRUCTURE DIAGRAM OF EACH MODULE", styles['BcaH2']))
    if os.path.exists('data/diagrams/activity_admin.png'):
        story.append(Image('data/diagrams/activity_admin.png', width=5*inch, height=3*inch))
    story.append(PageBreak())
    if os.path.exists('data/diagrams/activity_user.png'):
        story.append(Image('data/diagrams/activity_user.png', width=5*inch, height=4*inch))
        
    # Code
    story.append(PageBreak())
    story.append(Paragraph("4.1 PROGRAM LISTING", styles['BcaH2']))
    try:
        with open('src/pages/Index.tsx', 'r') as f:
            code = f.read()[:1000] + "\n... [Truncated]"
        story.append(Preformatted(code, styles['BcaCode']))
    except: pass

    # Screenshots
    story.append(PageBreak())
    story.append(Paragraph("4.2 INPUT SCREENS", styles['BcaH2']))
    if os.path.exists('public/screenshots/real_landing.png'):
        story.append(Image('public/screenshots/real_landing.png', width=6*inch, height=3.5*inch))
    
    story.append(PageBreak())
    story.append(Paragraph("4.3 OUTPUT SCREENS / REPORTS", styles['BcaH2']))
    if os.path.exists('public/screenshots/real_builder.png'):
        story.append(Image('public/screenshots/real_builder.png', width=6*inch, height=3.5*inch))
    story.append(Spacer(1, 0.2*inch))
    if os.path.exists('public/screenshots/real_dashboard.png'):
        story.append(Image('public/screenshots/real_dashboard.png', width=6*inch, height=3.5*inch))

    doc.build(story, onFirstPage=add_pdf_hf, onLaterPages=add_pdf_hf)

def main():
    os.makedirs('data/Generated_Reports', exist_ok=True)
    
    print("Generating DOCX files...")
    generate_docx('data/Generated_Reports/Student_Thesis_Report.docx')
    generate_docx('data/Generated_Reports/Practical_Project_Report.docx')
    
    print("Generating PDF files (to match exactly)...")
    generate_pdf('data/Generated_Reports/Student_Thesis_Report.pdf')
    generate_pdf('data/Generated_Reports/Practical_Project_Report.pdf')
    
    print("All files generated successfully in data/Generated_Reports/")

if __name__ == '__main__':
    main()
