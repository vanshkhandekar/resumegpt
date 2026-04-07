import os
import textwrap
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.colors import HexColor, black, white
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, Image, Preformatted, KeepTogether
)
from PIL import Image as PILImage

# --- CONFIGURATION ---
STUDENT_NAME = "Vansh Khandekar"
INSTITUTE = "Janaprabha Institute of Engineering and Technology, Ramtek"
COURSE = "Master of Computer Applications (MCA) Semester - I"
SUBJECT = "Advance Web Technologies"
BATCH = "2025 - 2026"
OUTPUT_FILE = "Practical_File.pdf"

THESIS_DIR = os.path.join(os.getcwd(), "thesis")
DIAG_DIR = os.path.join(THESIS_DIR, "thesis_diagrams")
SCREENSHOTS_DIR = os.path.join(os.getcwd(), "public/screenshots")

def get_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle('CoverTitle', fontName='Times-Bold', fontSize=28, leading=34, alignment=TA_CENTER, spaceAfter=20))
    styles.add(ParagraphStyle('CoverNormal', fontName='Times-Roman', fontSize=16, leading=22, alignment=TA_CENTER, spaceAfter=8))
    styles.add(ParagraphStyle('Heading', fontName='Times-Bold', fontSize=22, leading=26, alignment=TA_CENTER, spaceAfter=20))
    styles.add(ParagraphStyle('SubHeading', fontName='Times-Bold', fontSize=16, leading=20, alignment=TA_LEFT, spaceAfter=10))
    styles.add(ParagraphStyle('Body', fontName='Times-Roman', fontSize=12, leading=16, alignment=TA_JUSTIFY, spaceAfter=10))
    styles.add(ParagraphStyle('CodeBlock', fontName='Courier', fontSize=11, leading=14, spaceAfter=12, leftIndent=20, backColor=HexColor('#fafafa')))
    return styles

def spacer(h=12): return Spacer(1, h)

def get_code_chunks(file_path, start, end, chunk_size=45):
    """Reads a file and returns a list of wrapped code blocks."""
    try:
        if not os.path.exists(file_path): return [f"// Error: {file_path} not found"]
        with open(file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            snippet = lines[start-1:end]
            
            wrapped_snippet = []
            max_w = 60 # Safer width for 11pt Courier
            for line in snippet:
                clean_line = line.replace('\t', '    ').rstrip()
                if len(clean_line) > max_w:
                    indent = len(clean_line) - len(clean_line.lstrip())
                    wrapped = textwrap.fill(clean_line, width=max_w, subsequent_indent=' '*(indent+4))
                    wrapped_snippet.append(wrapped + '\n')
                else:
                    wrapped_snippet.append(clean_line + '\n')
                    
            chunks = []
            for i in range(0, len(wrapped_snippet), chunk_size):
                chunks.append("".join(wrapped_snippet[i:i+chunk_size]))
            return chunks
    except Exception as e: return [f"// Error: {str(e)}"]

def add_img(name_or_path, width=6.0*inch):
    path = name_or_path
    if not os.path.exists(path):
        path = os.path.join(SCREENSHOTS_DIR, name_or_path)
        if not os.path.exists(path):
             path = os.path.join(DIAG_DIR, name_or_path + ".png")
    if not os.path.exists(path): return None
    try:
        pil = PILImage.open(path)
        aspect = pil.size[1] / pil.size[0]
        height = width * aspect
        if height > 5.5 * inch:
            height = 5.5 * inch
            width = height / aspect
        return Image(path, width=width, height=height)
    except: return None

def build_cover(S):
    story = []
    story.append(spacer(100))
    story.append(Paragraph("PRACTICAL FILE", S['CoverTitle']))
    story.append(Paragraph(SUBJECT, S['CoverTitle']))
    story.append(spacer(50))
    story.append(Paragraph(f"<b>Student Name:</b> {STUDENT_NAME}", S['CoverNormal']))
    story.append(Paragraph(f"<b>Course:</b> {COURSE}", S['CoverNormal']))
    story.append(Paragraph(f"<b>{INSTITUTE}</b>", S['CoverNormal']))
    story.append(Paragraph(f"<b>Academic Year:</b> {BATCH}", S['CoverNormal']))
    story.append(PageBreak())
    return story

def main():
    doc = SimpleDocTemplate(OUTPUT_FILE, pagesize=A4, leftMargin=0.75*inch, rightMargin=0.75*inch)
    S = get_styles()
    story = []
    story += build_cover(S)

    practicals = [
        {"id": 1, "aim": "To implement a professional landing page using React and Framer Motion.", "code": ("src/pages/landing/LandingHero.tsx", 1, 500), "img": "real_landing.png"},
        {"id": 2, "aim": "To design a secure relational schema in PostgreSQL (Supabase) and define TypeScript interfaces.", "code": ("src/integrations/supabase/types.ts", 1, 600), "img": "detailed_er"},
        {"id": 3, "aim": "To develop a state-driven resume builder with dynamic section management.", "code": ("src/pages/dashboard/ResumeBuilder.tsx", 1, 700), "img": "real_builder.png"},
        {"id": 4, "aim": "To integrate Claude-3 Opus via OpenRouter API and implement system prompts for AI content generation.", "code": ("src/pages/Prompt.tsx", 1, 500), "img": "ai_sequence"},
        {"id": 5, "aim": "To implement a rule-based engine for evaluating keyword density and ATS scores.", "code": ("src/pages/dashboard/ResumeScore.tsx", 1, 600), "img": "real_dashboard.png"},
        {"id": 6, "aim": "To implement a browser-side PDF generation engine with high-fidelity layout rendering.", "code": ("src/pages/dashboard/ExportResume.tsx", 1, 600), "img": "real_templates.png"},
        {"id": 7, "aim": "To implement a custom React hook for debounced data persistence.", "code": ("src/hooks/useAutoSave.ts", 1, 500), "img": "state_user"},
        {"id": 8, "aim": "To create an administrative control panel for monitoring platform-wide metrics.", "code": ("src/pages/Admin.tsx", 1, 500), "img": "real_admin.png"},
        {"id": 9, "aim": "To implement secure JWT-based authentication using Supabase Auth.", "code": ("src/pages/Auth.tsx", 1, 500), "img": "auth_flow"},
        {"id": 10, "aim": "To design a scalable project architecture with global routing.", "code": ("src/App.tsx", 1, 400), "img": "architecture"},
        {"id": 11, "aim": "To implement custom business logic hooks for resume management.", "code": ("src/hooks/useResumes.ts", 1, 500), "img": "component_tree"},
        {"id": 12, "aim": "To implement a custom floating AI assistant with persistent chat history.", "code": ("src/components/ai/FloatingAiAssistant.tsx", 1, 600), "img": "sys_flow_user"},
        {"id": 13, "aim": "To implement responsive dashboard navigation.", "code": ("src/components/auth/UserMenu.tsx", 1, 400), "img": "sys_flow_admin"},
        {"id": 14, "aim": "To implement automated resume score improvements UI.", "code": ("src/components/dashboard/ScoreSuggestion.tsx", 1, 400), "img": "dfd_level1"}
    ]

    # Index
    story.append(Paragraph("INDEX", S['Heading']))
    index_data = [["Sr.No.", "Practical Aim", "Date", "Sign."]]
    for p in practicals:
        index_data.append([str(p['id']), p['aim'][:90] + "...", "", ""])
    t = Table(index_data, colWidths=[0.6*inch, 4.0*inch, 1*inch, 1*inch])
    t.setStyle(TableStyle([('GRID', (0,0), (-1,-1), 1, black), ('ALIGN', (0,0), (-1,-1), 'CENTER'), ('FONTNAME', (0,0), (-1,0), 'Times-Bold')]))
    story.append(t)
    story.append(PageBreak())

    # Practicals
    for p in practicals:
        story.append(Paragraph(f"Practical No. {p['id']}", S['SubHeading']))
        story.append(Paragraph(f"<b>Aim:</b> {p['aim']}", S['Body']))
        
        story.append(Paragraph("<b>Source Code Implementation:</b>", S['Body']))
        chunks = get_code_chunks(p['code'][0], p['code'][1], p['code'][2])
        for chunk in chunks:
            story.append(Preformatted(chunk, S['CodeBlock']))
        
        story.append(Paragraph("<b>Result Output:</b>", S['Body']))
        img_f = add_img(p['img'])
        if img_f:
             story.append(KeepTogether([img_f, Paragraph(f"<i>Figure {p['id']}: Practical System Output</i>", S['Body'])]))
        story.append(PageBreak())

    print(f"Generating PDF: {OUTPUT_FILE}")
    story_copy = list(story)
    doc.build(story)
    
    print(f"Generating DOCX: {OUTPUT_FILE.replace('.pdf', '.docx')}")
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        docx_doc = Document()
        for f in story_copy:
            if isinstance(f, Paragraph):
                text = f.getPlainText()
                para = docx_doc.add_paragraph()
                run = para.add_run(text)
                if 'Title' in f.style.name: run.bold=True; run.font.size=Pt(18)
                else: run.font.size=Pt(12)
            elif isinstance(f, Preformatted):
                para = docx_doc.add_paragraph()
                try: t = f.getPlainText()
                except: t = str(f)
                run = para.add_run(t)
                run.font.name = 'Courier New'
                run.font.size = Pt(8.5)
            elif isinstance(f, Image):
                try: docx_doc.add_picture(f.filename, width=Inches(6.0))
                except: pass
            elif isinstance(f, PageBreak): docx_doc.add_page_break()
        docx_doc.save(OUTPUT_FILE.replace('.pdf', '.docx'))
    except Exception as e: print(f"DOCX Failed: {e}")

if __name__ == "__main__":
    main()
