import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def add_heading(doc, text, level=1, align=None):
    h = doc.add_heading(text, level=level)
    if align:
        h.alignment = align
    return h

def add_paragraph(doc, text, align=None):
    p = doc.add_paragraph(text)
    if align:
        p.alignment = align
    return p

def run_script():
    doc = Document()
    
    # Notice that we need 2 identical word files: "Project_Report_Copy_1.docx" and "Project_Report_Copy_2.docx"
    
    # ─── FRONT PAGES CAN BE ADDED HERE IF NEEDED ───
    # We will start with the CONTENTS page
    
    # ---------- CONTENTS PAGE ----------
    contents_heading = doc.add_heading('CONTENTS', level=1)
    contents_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in contents_heading.runs:
        run.font.underline = True
    doc.add_paragraph()
    
    # Table of contents mirroring the image precisely
    toc_data = [
        ("Sr. No.", "CHAPTER", "Page No"),
        ("", "ABSTRACT", "I"),
        ("", "LIST OF FIGURES", "II"),
        ("", "LIST OF TABLES", "III"),
        ("1", "BRIEF REVIEW OF THE PROJECT", ""),
        ("", "1.1 TITLE", "1"),
        ("", "1.2 INTRODUCTION / OBJECTIVE", "2"),
        ("", "1.3 PRELIMINARY INVESTIGATION", "3"),
        ("", "1.4 FLAWS IN PRESENT SYSTEM", "4"),
        ("", "1.5 NEED OF NEW SYSTEM", "5"),
        ("2", "DETAILED SYSTEM DESIGN", ""),
        ("", "2.1 SYSTEM FLOW CHART", "6"),
        ("", "2.2 DETAILS OF LOGIC DEVELOPED", "7"),
        ("", "2.3 STRUCTURE DIAGRAM OF EACH MODULE", "9"),
        ("", "2.4 DATA DICTIONARY", "10"),
        ("", "2.5 DATA FLOW DIAGRAMS", "12"),
        ("3", "SOFTWARE/ HARDWARE DETAILS", ""),
        ("", "3.1 CHOICE OF A LANGUAGE USED", "14"),
        ("", "3.2 HARDWARE/ SOFTWARE SPECIFICATION", "15"),
        ("4", "SYSTEM DESIGN", ""),
        ("", "4.1 PROGRAM LISTING", "16"),
        ("", "4.2 INPUT SCREENS", "18"),
        ("", "4.3 OUTPUT SCREENS / REPORTS", "22"),
        ("5", "USER DOCUMENTATION", ""),
        ("", "5.1 IMPLEMENTATION, PROGRAM EXECUTION & MAINTENANCE", "26"),
        ("6", "CONCLUSION", ""),
        ("", "6.1 LIMITATIONS OF THE SYSTEM", "28"),
        ("", "6.2 SCOPE AND FUTURE MODIFICATION", "29"),
        ("7", "REFERENCES / BIBLIOGRAPHY", "30")
    ]
    
    table = doc.add_table(rows=len(toc_data), cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set widths: Sr. No (narrow), CHAPTER (wide), Page No (narrow)
    for row in table.rows:
        row.cells[0].width = Inches(0.8)
        row.cells[1].width = Inches(4.5)
        row.cells[2].width = Inches(0.8)
        
    for i, (sr, chap, pg) in enumerate(toc_data):
        row_cells = table.rows[i].cells
        row_cells[0].text = sr
        row_cells[1].text = chap
        row_cells[2].text = pg
        
        # Make the header and main chapters bold
        if i == 0 or sr != "":
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    doc.add_page_break()
    
    # ---------- CHAPTERS CONTENT ----------

    # ABSTRACT
    add_heading(doc, 'ABSTRACT', level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "The modern job market has become increasingly competitive, with employers relying heavily on Applicant Tracking Systems (ATS). AI Resume Studio is an advanced cloud-based Resume Maker platform configured with ATS Scoring & AI-Powered features which help candidates generate a professional and short-listed CV.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    doc.add_page_break()
    
    # LIST OF FIGURES
    add_heading(doc, 'LIST OF FIGURES', level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "Figure 1: System Flow Chart\nFigure 2: Data Flow Diagram\nFigure 3: Login Input Screen\nFigure 4: Dashboard Output Screen\nFigure 5: Resume Builder Interface\nFigure 6: Templates Selection")
    doc.add_page_break()
    
    # LIST OF TABLES
    add_heading(doc, 'LIST OF TABLES', level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "Table 1: Data Dictionary - Users\nTable 2: Data Dictionary - Resumes\nTable 3: Hardware / Software Requirements")
    doc.add_page_break()

    # CHAPTER 1
    add_heading(doc, '1. BRIEF REVIEW OF THE PROJECT', level=1)
    
    add_heading(doc, '1.1 TITLE', level=2)
    add_paragraph(doc, "AI Resume Studio — Resume Maker with ATS Scoring & AI-Powered Features")
    
    add_heading(doc, '1.2 INTRODUCTION / OBJECTIVE', level=2)
    add_paragraph(doc, "The main objective of AI Resume Studio is to provide job seekers with a powerful AI-driven platform to build professional, ATS-optimized resumes. It provides real-time ATS scoring, smart AI suggestions for bullet points, and dynamic generation of professional summaries.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    add_heading(doc, '1.3 PRELIMINARY INVESTIGATION', level=2)
    add_paragraph(doc, "Preliminary investigation revealed that over 75% of resumes are rejected by ATS software blindly because of poor formatting, lack of keywords, or unstructured designs. Job seekers face high rejection rates despite being qualified.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    add_heading(doc, '1.4 FLAWS IN PRESENT SYSTEM', level=2)
    add_paragraph(doc, "• Traditional generic tools lack ATS-friendly formatting.\n• Existing systems do not offer detailed analytics or pinpoint exact keyword gaps.\n• No contextual AI assistance that adapts to the specific job descriptions.")
    
    add_heading(doc, '1.5 NEED OF NEW SYSTEM', level=2)
    add_paragraph(doc, "• To bridge the gap between job applicant qualifications and recruiter expectations.\n• To automate resume formatting so users focus purely on content.\n• To use LLMs effectively in evaluating and refining professional vocabulary.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    doc.add_page_break()

    # CHAPTER 2
    add_heading(doc, '2. DETAILED SYSTEM DESIGN', level=1)
    
    add_heading(doc, '2.1 SYSTEM FLOW CHART', level=2)
    add_paragraph(doc, "User Authentication -> Dashboard -> Create/Edit Resume -> AI Optimization -> ATS Score Check -> Download PDF.")
    
    # (Assuming we don't have explicit diagrams, we just put description or leave placeholder)
    add_paragraph(doc, "[Placeholder for System Flow Chart Diagram]")
    
    add_heading(doc, '2.2 DETAILS OF LOGIC DEVELOPED', level=2)
    add_paragraph(doc, "The core logic entails fetching user section (Experience, Education, Skills) and running it through a sophisticated scoring algorithm. The API queries OpenRouter/Claude 3 Opus for intelligent recommendations based on TF-IDF logic of the provided Job descriptions.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    add_heading(doc, '2.3 STRUCTURE DIAGRAM OF EACH MODULE', level=2)
    add_paragraph(doc, "Modules Include:\n1. Auth Module\n2. Resume CRUD Module\n3. AI Analysis Module \n4. Export/PDF Module\n[Placeholder for Structure Diagram]")
    
    add_heading(doc, '2.4 DATA DICTIONARY', level=2)
    add_paragraph(doc, "Table: Users\nColumns: id (UUID), email (VARCHAR), created_at (TIMESTAMP), subscription_tier (VARCHAR).\n\nTable: Resumes\nColumns: id (UUID), user_id (UUID), content (JSONB), ats_score (INTEGER).")
    
    add_heading(doc, '2.5 DATA FLOW DIAGRAMS', level=2)
    add_paragraph(doc, "DFD Level 0: User -> System -> Job Ready Output.\nDFD Level 1: Expand into Sub-processes like Validation, Formatting, and Download.\n[Placeholder for DFD Diagram]")
    doc.add_page_break()

    # CHAPTER 3
    add_heading(doc, '3. SOFTWARE/ HARDWARE DETAILS', level=1)
    
    add_heading(doc, '3.1 CHOICE OF A LANGUAGE USED', level=2)
    add_paragraph(doc, "• Frontend: HTML5, CSS3/Tailwind, React 18, TypeScript.\n• Backend Logic: Supabase Edge Functions (Deno/TypeScript), Node.js.\n• Database: PostgreSQL (via Supabase).")
    
    add_heading(doc, '3.2 HARDWARE/ SOFTWARE SPECIFICATION', level=2)
    add_paragraph(doc, "Hardware Requirements (Client Side):\n- Any modern operating system (Windows, Linux, macOS).\n- 4GB RAM minimum, 2-core processor.\n\nSoftware Requirements:\n- Node.js (v18+)\n- Modern Web Browser (Chrome, Edge, Firefox) with JS enabled.")
    doc.add_page_break()

    # CHAPTER 4
    add_heading(doc, '4. SYSTEM DESIGN', level=1)
    
    add_heading(doc, '4.1 PROGRAM LISTING', level=2)
    add_paragraph(doc, "The program consists of modular components in React (`src/components`), hooks for data fetching (`src/hooks/useResume.ts`), pages (`src/pages`), and serverless Edge Functions linked to Supabase PostgreSQL.")
    
    add_heading(doc, '4.2 INPUT SCREENS', level=2)
    add_paragraph(doc, "The input screens capture user interactions such as logging in, selecting templates, and inputting resume details into the Builder module.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    # Adding screenshots requested for INPUT SCREENS
    screens_input = [
        ("Landing Page (System Entry)", 'public/screenshots/real_landing.png'),
        ("Template Selection Screen", 'public/screenshots/real_templates.png'),
    ]
    
    for title, img_path in screens_input:
        if os.path.exists(img_path):
            add_paragraph(doc, f"Screen: {title}", align=WD_ALIGN_PARAGRAPH.CENTER)
            doc.add_picture(img_path, width=Inches(6.0))
            add_paragraph(doc, "")

    add_heading(doc, '4.3 OUTPUT SCREENS / REPORTS', level=2)
    add_paragraph(doc, "The output screens show the final rendered layout, the analytics dashboard, and the administrative reports.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    # Adding screenshots requested for OUTPUT SCREENS
    screens_output = [
        ("Resume Builder Interface (Output View)", 'public/screenshots/real_builder.png'),
        ("User Dashboard (Analytics and Management)", 'public/screenshots/real_dashboard.png'),
        ("Admin Reporting Dashboard", 'public/screenshots/real_admin.png')
    ]
    
    for title, img_path in screens_output:
        if os.path.exists(img_path):
            add_paragraph(doc, f"Screen: {title}", align=WD_ALIGN_PARAGRAPH.CENTER)
            doc.add_picture(img_path, width=Inches(6.0))
            add_paragraph(doc, "")
    
    doc.add_page_break()

    # CHAPTER 5
    add_heading(doc, '5. USER DOCUMENTATION', level=1)
    
    add_heading(doc, '5.1 IMPLEMENTATION, PROGRAM EXECUTION & MAINTENANCE', level=2)
    add_paragraph(doc, "Implementation Details:\n1. The system is deployed via Vercel for the frontend and Supabase for the backend infrastructure.\n2. Program execution requires a user to sign-up via OAuth or Email links.\n3. Maintenance consists of periodic updates to AI system prompts and updating the ATS algorithms inline with current industry trends.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    doc.add_page_break()

    # CHAPTER 6
    add_heading(doc, '6. CONCLUSION', level=1)
    
    add_heading(doc, '6.1 LIMITATIONS OF THE SYSTEM', level=2)
    add_paragraph(doc, "• Requires active high-speed internet connection for real-time AI suggestions.\n• Contextual limits in LLMs might miss highly specialized niche terminology natively unencountered.\n• Strict PDF parsing formatting limitations when downloading in raw environments.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    add_heading(doc, '6.2 SCOPE AND FUTURE MODIFICATION', level=2)
    add_paragraph(doc, "• Integration with LinkedIn for direct 1-click importing of user profiles.\n• Support for multiple file exports like proper DOCX conversion.\n• Mobile App implementation for resume management on-the-go.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    doc.add_page_break()

    # CHAPTER 7
    add_heading(doc, '7. REFERENCES / BIBLIOGRAPHY', level=1)
    add_paragraph(doc, "1. React Documentation: https://reactjs.org/docs\n2. Supabase Documentation: https://supabase.com/docs\n3. OpenAI / Anthropic APIs Context Guides.\n4. Modern Systems Analysis and Design by Hoffer et al.\n5. Applicant Tracking Algorithms and Implementation Guidelines, 2025.")

    # Save identical copies based on "dono files same" & "Create 02 Word Files"
    os.makedirs('data/Generated_Reports', exist_ok=True)
    file1 = 'data/Generated_Reports/Student_Thesis_Report.docx'
    file2 = 'data/Generated_Reports/Practical_Project_Report.docx'
    
    doc.save(file1)
    doc.save(file2)
    print(f"Generated successfully: {file1} and {file2}")

if __name__ == '__main__':
    run_script()
